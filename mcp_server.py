#!/usr/bin/env python3
"""MCP Server for document operations (docx, xlsx, pdf, txt, csv)"""

import json
import os
from typing import Optional, List, Any, Dict
from mcp.server import Server
from mcp.server.stdio import stdio_server
import mcp.types as types

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
import pdfplumber
import pandas as pd

app = Server("doc-tools")


def _validate_path(path: str) -> str:
    abs_path = os.path.abspath(path)
    if not os.path.exists(abs_path):
        raise FileNotFoundError(f"File not found: {abs_path}")
    return abs_path


@app.list_tools()
async def list_tools() -> List[types.Tool]:
    return [
        types.Tool(
            name="read_txt",
            description="Read a text file and return its contents",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the text file"}
                },
                "required": ["path"]
            }
        ),
        types.Tool(
            name="read_csv",
            description="Read a CSV file and return its contents as JSON",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the CSV file"},
                    "delimiter": {"type": "string", "default": ",", "description": "CSV delimiter"}
                },
                "required": ["path"]
            }
        ),
        types.Tool(
            name="read_docx",
            description="Read a Word document and return its text content",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the .docx file"}
                },
                "required": ["path"]
            }
        ),
        types.Tool(
            name="read_xlsx",
            description="Read an Excel file and return its contents as JSON (all sheets)",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the .xlsx file"},
                    "sheet": {"type": "string", "description": "Specific sheet name (optional, returns all if not specified)"}
                },
                "required": ["path"]
            }
        ),
        types.Tool(
            name="read_pdf",
            description="Read a PDF file and return its text content",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the PDF file"},
                    "pages": {"type": "string", "description": "Page range (e.g., '1-5' or '1,3,5'), optional"}
                },
                "required": ["path"]
            }
        ),
        types.Tool(
            name="create_txt",
            description="Create a text file with the given content",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path for the new text file"},
                    "content": {"type": "string", "description": "Content to write"}
                },
                "required": ["path", "content"]
            }
        ),
        types.Tool(
            name="create_csv",
            description="Create a CSV file from JSON data",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path for the new CSV file"},
                    "data": {"type": "string", "description": "JSON array of objects to write as CSV"},
                    "delimiter": {"type": "string", "default": ",", "description": "CSV delimiter"}
                },
                "required": ["path", "data"]
            }
        ),
        types.Tool(
            name="create_docx",
            description="Create a Word document with text content",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path for the new .docx file"},
                    "content": {"type": "string", "description": "Text content (use \\n for paragraphs)"}
                },
                "required": ["path", "content"]
            }
        ),
        types.Tool(
            name="create_xlsx",
            description="Create an Excel file from JSON data",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path for the new .xlsx file"},
                    "sheets": {"type": "string", "description": "JSON object: {\"SheetName\": [[row1], [row2], ...] or [{col: val}, ...]}"},
                    "headers": {"type": "boolean", "default": True, "description": "First row is header (for object data)"}
                },
                "required": ["path", "sheets"]
            }
        ),
        types.Tool(
            name="append_to_txt",
            description="Append content to an existing text file",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the text file"},
                    "content": {"type": "string", "description": "Content to append"}
                },
                "required": ["path", "content"]
            }
        ),
        types.Tool(
            name="append_to_csv",
            description="Append rows to an existing CSV file",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to the CSV file"},
                    "data": {"type": "string", "description": "JSON array of objects to append"}
                },
                "required": ["path", "data"]
            }
        ),
        types.Tool(
            name="list_directory",
            description="List files in a directory",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Directory path"},
                    "pattern": {"type": "string", "description": "File pattern filter (e.g., '*.txt'), optional"}
                },
                "required": ["path"]
            }
        ),
        types.Tool(
            name="get_file_info",
            description="Get file metadata (size, modified time, etc.)",
            inputSchema={
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Path to file"}
                },
                "required": ["path"]
            }
        )
    ]


@app.call_tool()
async def call_tool(name: str, arguments: Dict[str, Any]) -> List[types.TextContent]:
    try:
        if name == "read_txt":
            path = _validate_path(arguments["path"])
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()
            return [types.TextContent(type="text", text=content)]

        elif name == "read_csv":
            path = _validate_path(arguments["path"])
            delimiter = arguments.get("delimiter", ",")
            df = pd.read_csv(path, delimiter=delimiter)
            result = df.to_json(orient='records', force_ascii=False)
            return [types.TextContent(type="text", text=result)]

        elif name == "read_docx":
            path = _validate_path(arguments["path"])
            doc = Document(path)
            content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            for table in doc.tables:
                content.append("\n[TABLE]")
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    content.append(row_text)
            return [types.TextContent(type="text", text="\n".join(content))]

        elif name == "read_xlsx":
            path = _validate_path(arguments["path"])
            sheet_name = arguments.get("sheet")
            wb = load_workbook(path, data_only=True)
            
            result = {}
            sheets_to_read = [sheet_name] if sheet_name else wb.sheetnames
            
            for sheet in sheets_to_read:
                if sheet not in wb.sheetnames:
                    continue
                ws = wb[sheet]
                data = []
                for row in ws.iter_rows(values_only=True):
                    data.append([str(cell) if cell is not None else "" for cell in row])
                result[sheet] = data
            
            return [types.TextContent(type="text", text=json.dumps(result, ensure_ascii=False))]

        elif name == "read_pdf":
            path = _validate_path(arguments["path"])
            pages_str = arguments.get("pages")
            
            text_content = []
            with pdfplumber.open(path) as pdf:
                if pages_str:
                    pages = _parse_page_range(pages_str, len(pdf.pages))
                else:
                    pages = range(len(pdf.pages))
                
                for i in pages:
                    page = pdf.pages[i]
                    text = page.extract_text() or ""
                    text_content.append(f"--- Page {i+1} ---\n{text}")
            
            return [types.TextContent(type="text", text="\n\n".join(text_content))]

        elif name == "create_txt":
            path = arguments["path"]
            content = arguments["content"]
            os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
            with open(path, 'w', encoding='utf-8') as f:
                f.write(content)
            return [types.TextContent(type="text", text=f"Created: {path}")]

        elif name == "create_csv":
            path = arguments["path"]
            data = json.loads(arguments["data"])
            delimiter = arguments.get("delimiter", ",")
            os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
            df = pd.DataFrame(data)
            df.to_csv(path, index=False, sep=delimiter)
            return [types.TextContent(type="text", text=f"Created: {path} ({len(data)} rows)")]

        elif name == "create_docx":
            path = arguments["path"]
            content = arguments["content"]
            os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
            doc = Document()
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
            doc.save(path)
            return [types.TextContent(type="text", text=f"Created: {path}")]

        elif name == "create_xlsx":
            path = arguments["path"]
            sheets_data = json.loads(arguments["sheets"])
            use_headers = arguments.get("headers", True)
            os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
            
            wb = Workbook()
            first_sheet = True
            
            for sheet_name, data in sheets_data.items():
                if first_sheet:
                    ws = wb.active
                    ws.title = sheet_name
                    first_sheet = False
                else:
                    ws = wb.create_sheet(sheet_name)
                
                if data and isinstance(data[0], dict):
                    headers = list(data[0].keys())
                    if use_headers:
                        ws.append(headers)
                    for row_data in data:
                        ws.append([row_data.get(h, "") for h in headers])
                else:
                    for row in data:
                        ws.append(row)
            
            wb.save(path)
            return [types.TextContent(type="text", text=f"Created: {path}")]

        elif name == "append_to_txt":
            path = arguments["path"]
            content = arguments["content"]
            with open(path, 'a', encoding='utf-8') as f:
                f.write("\n" + content)
            return [types.TextContent(type="text", text=f"Appended to: {path}")]

        elif name == "append_to_csv":
            path = _validate_path(arguments["path"])
            data = json.loads(arguments["data"])
            df = pd.DataFrame(data)
            df.to_csv(path, mode='a', index=False, header=False)
            return [types.TextContent(type="text", text=f"Appended {len(data)} rows to: {path}")]

        elif name == "list_directory":
            path = arguments["path"]
            pattern = arguments.get("pattern", "*")
            import glob
            files = glob.glob(os.path.join(path, pattern))
            result = [{"name": os.path.basename(f), "path": f, "is_dir": os.path.isdir(f)} for f in files]
            return [types.TextContent(type="text", text=json.dumps(result, ensure_ascii=False, indent=2))]

        elif name == "get_file_info":
            path = _validate_path(arguments["path"])
            stat = os.stat(path)
            info = {
                "path": path,
                "size": stat.st_size,
                "size_human": _human_readable_size(stat.st_size),
                "modified": stat.st_mtime,
                "is_file": os.path.isfile(path),
                "is_dir": os.path.isdir(path)
            }
            return [types.TextContent(type="text", text=json.dumps(info, ensure_ascii=False, indent=2))]

        else:
            return [types.TextContent(type="text", text=f"Unknown tool: {name}")]

    except Exception as e:
        return [types.TextContent(type="text", text=f"Error: {str(e)}")]


def _parse_page_range(pages_str: str, total_pages: int) -> List[int]:
    pages = set()
    for part in pages_str.split(','):
        part = part.strip()
        if '-' in part:
            start, end = map(int, part.split('-'))
            pages.update(range(start - 1, min(end, total_pages)))
        else:
            page = int(part) - 1
            if 0 <= page < total_pages:
                pages.add(page)
    return sorted(pages)


def _human_readable_size(size: int) -> str:
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"


async def main():
    async with stdio_server() as (read_stream, write_stream):
        await app.run(read_stream, write_stream, app.create_initialization_options())


if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
