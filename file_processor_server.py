#!/usr/bin/env python3
"""
Unified file processing MCP server
Upload files to /files/uploads/ and process them automatically
"""

import os
import sys
import json
import hashlib
import shutil
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional

from mcp.server import Server
from mcp.server.stdio import stdio_server
import mcp.types as types

UPLOADS_DIR = "/files/uploads"
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
ALLOWED_EXTENSIONS = {
    'documents': ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv', '.txt', '.rtf', '.odt'],
    'images': ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg'],
    'data': ['.json', '.xml', '.yaml', '.yml'],
    'archives': ['.zip', '.tar', '.gz', '.rar'],
    'code': ['.py', '.js', '.ts', '.html', '.css', '.sql', '.sh', '.java', '.cpp', '.c', '.go', '.rs']
}
ALL_ALLOWED = set()
for exts in ALLOWED_EXTENSIONS.values():
    ALL_ALLOWED.update(exts)

os.makedirs(UPLOADS_DIR, exist_ok=True)

app = Server("file-processor")


def get_file_hash(filepath: str) -> str:
    h = hashlib.md5()
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            h.update(chunk)
    return h.hexdigest()


def get_file_category(ext: str) -> str:
    ext = ext.lower()
    for category, exts in ALLOWED_EXTENSIONS.items():
        if ext in exts:
            return category
    return 'unknown'


def cleanup_old_files(hours: int = 24):
    """Remove files older than specified hours"""
    removed = 0
    cutoff = datetime.now() - timedelta(hours=hours)
    for f in os.listdir(UPLOADS_DIR):
        filepath = os.path.join(UPLOADS_DIR, f)
        if os.path.isfile(filepath):
            mtime = datetime.fromtimestamp(os.path.getmtime(filepath))
            if mtime < cutoff:
                os.remove(filepath)
                removed += 1
    return removed


def process_pdf(filepath: str) -> Dict:
    """Extract text from PDF"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                text_parts.append(f"--- Page {i+1} ---\n{text}")
        return {
            'success': True,
            'type': 'pdf',
            'pages': len(pdf.pages) if 'pdf' in locals() else 0,
            'content': '\n\n'.join(text_parts)
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}


def process_docx(filepath: str) -> Dict:
    """Extract text from DOCX"""
    try:
        from docx import Document
        doc = Document(filepath)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        for table in doc.tables:
            content.append("\n[TABLE]")
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                content.append(row_text)
        return {
            'success': True,
            'type': 'docx',
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'content': '\n'.join(content)
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}


def process_xlsx(filepath: str, max_rows: int = 1000) -> Dict:
    """Extract data from XLSX"""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(filepath, data_only=True)
        result = {'success': True, 'type': 'xlsx', 'sheets': {}}
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            data = []
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= max_rows:
                    data.append(['... truncated ...'])
                    break
                data.append([str(cell) if cell is not None else "" for cell in row])
                row_count += 1
            result['sheets'][sheet_name] = data
        return result
    except Exception as e:
        return {'success': False, 'error': str(e)}


def process_csv(filepath: str, max_rows: int = 1000) -> Dict:
    """Extract data from CSV"""
    try:
        import pandas as pd
        df = pd.read_csv(filepath, nrows=max_rows)
        return {
            'success': True,
            'type': 'csv',
            'rows': len(df),
            'columns': list(df.columns),
            'content': df.to_json(orient='records', force_ascii=False)
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}


def process_txt(filepath: str, max_chars: int = 50000) -> Dict:
    """Read text file"""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read(max_chars)
            truncated = len(content) == max_chars
        return {
            'success': True,
            'type': 'text',
            'truncated': truncated,
            'content': content
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}


def process_json(filepath: str) -> Dict:
    """Read JSON file"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        return {
            'success': True,
            'type': 'json',
            'content': json.dumps(data, ensure_ascii=False, indent=2)
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}


def process_image(filepath: str) -> Dict:
    """Get image metadata"""
    try:
        from PIL import Image
        with Image.open(filepath) as img:
            return {
                'success': True,
                'type': 'image',
                'format': img.format,
                'size': img.size,
                'mode': img.mode,
                'info': f"Image: {img.size[0]}x{img.size[1]}, {img.format}, {img.mode}"
            }
    except Exception as e:
        return {'success': False, 'error': str(e)}


def process_archive(filepath: str) -> Dict:
    """List archive contents"""
    try:
        import zipfile
        import tarfile
        
        ext = os.path.splitext(filepath)[1].lower()
        files = []
        
        if ext == '.zip':
            with zipfile.ZipFile(filepath, 'r') as z:
                files = z.namelist()
        elif ext in ['.tar', '.gz', '.tgz']:
            with tarfile.open(filepath, 'r:*') as t:
                files = [m.name for m in t.getmembers()]
        
        return {
            'success': True,
            'type': 'archive',
            'files_count': len(files),
            'files': files[:100] if len(files) > 100 else files,
            'truncated': len(files) > 100
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}


def process_file(filepath: str) -> Dict:
    """Auto-detect file type and process"""
    if not os.path.exists(filepath):
        return {'success': False, 'error': 'File not found'}
    
    file_size = os.path.getsize(filepath)
    if file_size > MAX_FILE_SIZE:
        return {'success': False, 'error': f'File too large: {file_size} bytes (max: {MAX_FILE_SIZE})'}
    
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.pdf':
        return process_pdf(filepath)
    elif ext in ['.docx', '.doc']:
        return process_docx(filepath)
    elif ext in ['.xlsx', '.xls']:
        return process_xlsx(filepath)
    elif ext == '.csv':
        return process_csv(filepath)
    elif ext in ['.txt', '.rtf', '.md', '.log']:
        return process_txt(filepath)
    elif ext == '.json':
        return process_json(filepath)
    elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp']:
        return process_image(filepath)
    elif ext in ['.zip', '.tar', '.gz', '.tgz', '.rar']:
        return process_archive(filepath)
    elif ext in ALL_ALLOWED:
        return process_txt(filepath)
    else:
        return {'success': False, 'error': f'Unsupported file type: {ext}'}


@app.list_tools()
async def list_tools() -> List[types.Tool]:
    return [
        types.Tool(
            name="scan_uploads",
            description="Scan uploads directory and list all uploaded files with metadata",
            inputSchema={
                "type": "object",
                "properties": {
                    "cleanup_hours": {"type": "integer", "description": "Remove files older than N hours (0 = no cleanup)"}
                },
                "required": []
            }
        ),
        types.Tool(
            name="process_uploaded_file",
            description="Process an uploaded file (auto-detect type: PDF, DOCX, XLSX, CSV, TXT, JSON, images, archives)",
            inputSchema={
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Name of the file in uploads directory"},
                    "action": {"type": "string", "enum": ["read", "info", "delete"], "description": "Action: read (extract content), info (metadata only), delete"}
                },
                "required": ["filename"]
            }
        ),
        types.Tool(
            name="process_all_uploads",
            description="Process all files in uploads directory",
            inputSchema={
                "type": "object",
                "properties": {
                    "cleanup_hours": {"type": "integer", "description": "Remove files older than N hours after processing (0 = no cleanup)"}
                },
                "required": []
            }
        ),
        types.Tool(
            name="get_upload_path",
            description="Get the path where user should upload files",
            inputSchema={
                "type": "object",
                "properties": {},
                "required": []
            }
        )
    ]


@app.call_tool()
async def call_tool(name: str, arguments: Dict[str, Any]) -> List[types.TextContent]:
    try:
        if name == "get_upload_path":
            return [types.TextContent(
                type="text",
                text=f"Upload files to:\n\n{UPLOADS_DIR}/\n\nSupported formats:\n"
                     f"- Documents: PDF, DOCX, XLSX, CSV, TXT\n"
                     f"- Images: PNG, JPG, GIF, BMP\n"
                     f"- Data: JSON, XML, YAML\n"
                     f"- Archives: ZIP, TAR, GZ\n"
                     f"- Code: PY, JS, HTML, CSS, SQL\n\n"
                     f"Max file size: {MAX_FILE_SIZE // (1024*1024)} MB"
            )]
        
        elif name == "scan_uploads":
            cleanup_hours = arguments.get("cleanup_hours", 0)
            if cleanup_hours > 0:
                removed = cleanup_old_files(cleanup_hours)
            
            files = []
            for f in os.listdir(UPLOADS_DIR):
                filepath = os.path.join(UPLOADS_DIR, f)
                if os.path.isfile(filepath):
                    stat = os.stat(filepath)
                    ext = os.path.splitext(f)[1].lower()
                    files.append({
                        'name': f,
                        'size': stat.st_size,
                        'size_human': f"{stat.st_size / 1024:.1f} KB" if stat.st_size < 1024*1024 else f"{stat.st_size / (1024*1024):.1f} MB",
                        'modified': datetime.fromtimestamp(stat.st_mtime).isoformat()[:19],
                        'category': get_file_category(ext),
                        'extension': ext
                    })
            
            if not files:
                return [types.TextContent(type="text", text=f"No files in uploads directory.\n\nUpload to: {UPLOADS_DIR}/")]
            
            files.sort(key=lambda x: x['modified'], reverse=True)
            
            result = f"Uploaded files ({len(files)}):\n\n"
            for f in files:
                result += f"• {f['name']}\n"
                result += f"  Size: {f['size_human']} | Type: {f['category']} | Modified: {f['modified']}\n\n"
            
            if cleanup_hours > 0:
                result += f"\nCleaned up {removed} old files."
            
            return [types.TextContent(type="text", text=result)]
        
        elif name == "process_uploaded_file":
            filename = arguments["filename"]
            action = arguments.get("action", "read")
            
            filepath = os.path.join(UPLOADS_DIR, filename)
            
            if not os.path.exists(filepath):
                return [types.TextContent(type="text", text=f"File not found: {filename}\n\nUpload to: {UPLOADS_DIR}/")]
            
            if action == "delete":
                os.remove(filepath)
                return [types.TextContent(type="text", text=f"Deleted: {filename}")]
            
            if action == "info":
                stat = os.stat(filepath)
                ext = os.path.splitext(filename)[1].lower()
                info = {
                    'filename': filename,
                    'path': filepath,
                    'size': stat.st_size,
                    'size_human': f"{stat.st_size / 1024:.1f} KB" if stat.st_size < 1024*1024 else f"{stat.st_size / (1024*1024):.1f} MB",
                    'modified': datetime.fromtimestamp(stat.st_mtime).isoformat()[:19],
                    'category': get_file_category(ext),
                    'hash': get_file_hash(filepath)
                }
                return [types.TextContent(type="text", text=json.dumps(info, indent=2))]
            
            result = process_file(filepath)
            
            if result['success']:
                if result['type'] in ['pdf', 'docx', 'text', 'json']:
                    content = result.get('content', '')
                    if len(content) > 30000:
                        content = content[:30000] + "\n\n... [content truncated]"
                    return [types.TextContent(type="text", text=f"File: {filename}\nType: {result['type']}\n\n{content}")]
                elif result['type'] == 'xlsx':
                    return [types.TextContent(type="text", text=f"File: {filename}\nType: XLSX\n\n{json.dumps(result['sheets'], ensure_ascii=False, indent=2)}")]
                elif result['type'] == 'csv':
                    return [types.TextContent(type="text", text=f"File: {filename}\nType: CSV\nRows: {result['rows']}\nColumns: {result['columns']}\n\n{result['content'][:10000]}")]
                elif result['type'] == 'image':
                    return [types.TextContent(type="text", text=f"File: {filename}\nType: Image\n{result['info']}")]
                elif result['type'] == 'archive':
                    files_list = '\n'.join(result['files'][:50])
                    return [types.TextContent(type="text", text=f"File: {filename}\nType: Archive\nFiles: {result['files_count']}\n\n{files_list}")]
                else:
                    return [types.TextContent(type="text", text=f"File: {filename}\n{json.dumps(result, indent=2)}")]
            else:
                return [types.TextContent(type="text", text=f"Error processing {filename}: {result['error']}")]
        
        elif name == "process_all_uploads":
            cleanup_hours = arguments.get("cleanup_hours", 0)
            
            files = [f for f in os.listdir(UPLOADS_DIR) if os.path.isfile(os.path.join(UPLOADS_DIR, f))]
            
            if not files:
                return [types.TextContent(type="text", text=f"No files to process.\n\nUpload to: {UPLOADS_DIR}/")]
            
            results = []
            for filename in files:
                filepath = os.path.join(UPLOADS_DIR, filename)
                result = process_file(filepath)
                results.append({'filename': filename, 'success': result['success'], 'type': result.get('type', 'unknown'), 'error': result.get('error')})
            
            if cleanup_hours > 0:
                removed = cleanup_old_files(cleanup_hours)
            
            summary = f"Processed {len(results)} files:\n\n"
            for r in results:
                status = "✓" if r['success'] else "✗"
                summary += f"{status} {r['filename']} ({r['type']})\n"
                if not r['success']:
                    summary += f"  Error: {r['error']}\n"
            
            if cleanup_hours > 0:
                summary += f"\nCleaned up {removed} old files."
            
            return [types.TextContent(type="text", text=summary)]
        
        else:
            return [types.TextContent(type="text", text=f"Unknown tool: {name}")]
    
    except Exception as e:
        return [types.TextContent(type="text", text=f"Error: {str(e)}")]


async def main():
    async with stdio_server() as (read_stream, write_stream):
        await app.run(read_stream, write_stream, app.create_initialization_options())


if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
