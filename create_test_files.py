#!/usr/bin/env python3
"""Create test documents for mcpdocs with demo data"""

import os
import random
from datetime import datetime, timedelta
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors

TEST_DIR = "/home/projects/mcpdocs/test_files"
os.makedirs(TEST_DIR, exist_ok=True)

# Create test DOCX - Business Report
doc = Document()
doc.add_heading('Business Report Q4 2025', 0)
doc.add_paragraph('CONFIDENTIAL - Internal Use Only')
doc.add_paragraph('Generated: January 15, 2026')

doc.add_heading('Executive Summary', level=1)
doc.add_paragraph('This report provides a comprehensive overview of the company\'s performance during Q4 2025. The quarter showed strong growth in revenue and customer acquisition.')

doc.add_heading('Key Metrics', level=1)
metrics = [
    ('Total Revenue', '$2,450,000', '+15% YoY'),
    ('New Customers', '1,247', '+22% YoY'),
    ('Churn Rate', '2.1%', '-0.5% YoY'),
    ('Net Promoter Score', '72', '+5 pts YoY'),
]
for metric, value, change in metrics:
    p = doc.add_paragraph()
    p.add_run(f'{metric}: ').bold = True
    p.add_run(f'{value} ')
    p.add_run(f'({change})').italic = True

doc.add_heading('Department Performance', level=1)
doc.add_paragraph('Sales Department exceeded targets by 18%. Engineering delivered 12 major features. Marketing campaign ROI increased by 25%.')

doc.add_heading('Challenges & Recommendations', level=1)
doc.add_paragraph('1. Supply chain delays - recommend diversifying suppliers')
doc.add_paragraph('2. Talent retention - implement new compensation packages')
doc.add_paragraph('3. Technical debt - allocate 20% of sprint capacity')

table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
headers = ['Region', 'Revenue', 'Growth', 'Target']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
data = [
    ['North America', '$980,000', '+12%', '$900,000'],
    ['Europe', '$720,000', '+18%', '$650,000'],
    ['Asia Pacific', '$520,000', '+25%', '$450,000'],
    ['Latin America', '$230,000', '+8%', '$250,000'],
]
for row_idx, row_data in enumerate(data, 1):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = cell_data

doc.add_paragraph()
doc.add_paragraph('Prepared by: Analytics Team')
doc.add_paragraph('Approved by: CFO')
doc.save(os.path.join(TEST_DIR, 'business_report.docx'))
print("Created: business_report.docx")

# Create test XLSX - Financial Report + Employees + Sales
wb = Workbook()

# Sheet 1: Financial Summary
ws = wb.active
ws.title = "Financial Summary"
ws.merge_cells('A1:F1')
ws['A1'] = 'Financial Summary 2025'
ws['A1'].font = Font(size=16, bold=True)
ws['A1'].alignment = Alignment(horizontal='center')

headers = ['Month', 'Revenue', 'Expenses', 'Profit', 'Margin %', 'YoY Growth']
ws.append([])
for col, header in enumerate(headers, 1):
    cell = ws.cell(row=3, column=col, value=header)
    cell.font = Font(bold=True)
    cell.fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    cell.font = Font(bold=True, color='FFFFFF')

financial_data = [
    ['January', 185000, 142000, 43000, 23.2, 12],
    ['February', 192000, 148000, 44000, 22.9, 14],
    ['March', 198000, 151000, 47000, 23.7, 11],
    ['April', 205000, 155000, 50000, 24.4, 16],
    ['May', 218000, 162000, 56000, 25.7, 18],
    ['June', 225000, 168000, 57000, 25.3, 15],
    ['July', 232000, 172000, 60000, 25.9, 17],
    ['August', 241000, 178000, 63000, 26.1, 19],
    ['September', 248000, 182000, 66000, 26.6, 21],
    ['October', 255000, 186000, 69000, 27.1, 18],
    ['November', 268000, 192000, 76000, 28.4, 22],
    ['December', 283000, 198000, 85000, 30.0, 25],
]
for row in financial_data:
    ws.append(row)

ws.append([])
ws.append(['Total', '=SUM(B4:B15)', '=SUM(C4:C15)', '=SUM(D4:D15)', '=AVERAGE(E4:E15)', '=AVERAGE(F4:F15)'])
ws.append(['Q1 Total', '=SUM(B4:B6)', '=SUM(C4:C6)', '=SUM(D4:D6)', '', ''])
ws.append(['Q2 Total', '=SUM(B7:B9)', '=SUM(C7:C9)', '=SUM(D7:D9)', '', ''])
ws.append(['Q3 Total', '=SUM(B10:B12)', '=SUM(C10:C12)', '=SUM(D10:D12)', '', ''])
ws.append(['Q4 Total', '=SUM(B13:B15)', '=SUM(C13:C15)', '=SUM(D13:D15)', '', ''])

for col in range(1, 7):
    ws.column_dimensions[get_column_letter(col)].width = 14

# Sheet 2: Employees
ws2 = wb.create_sheet("Employees")
headers = ['ID', 'Name', 'Department', 'Position', 'Salary', 'Start Date', 'Email']
ws2.append(headers)
for col in range(1, 8):
    ws2.cell(row=1, column=col).font = Font(bold=True)
    ws2.cell(row=1, column=col).fill = PatternFill(start_color='70AD47', end_color='70AD47', fill_type='solid')
    ws2.cell(row=1, column=col).font = Font(bold=True, color='FFFFFF')

employees = [
    [1, 'John Smith', 'Engineering', 'Senior Developer', 95000, '2022-03-15', 'john.smith@company.com'],
    [2, 'Sarah Johnson', 'Marketing', 'Marketing Manager', 85000, '2021-08-01', 'sarah.j@company.com'],
    [3, 'Michael Chen', 'Sales', 'Sales Director', 110000, '2020-01-10', 'm.chen@company.com'],
    [4, 'Emily Davis', 'HR', 'HR Specialist', 65000, '2023-02-20', 'emily.d@company.com'],
    [5, 'Robert Wilson', 'Engineering', 'DevOps Engineer', 90000, '2022-11-05', 'r.wilson@company.com'],
    [6, 'Lisa Anderson', 'Finance', 'Financial Analyst', 75000, '2021-06-15', 'lisa.a@company.com'],
    [7, 'James Taylor', 'Operations', 'Operations Manager', 80000, '2020-09-01', 'james.t@company.com'],
    [8, 'Jennifer Lee', 'Engineering', 'Frontend Developer', 72000, '2023-05-10', 'jennifer.l@company.com'],
    [9, 'David Brown', 'Sales', 'Account Executive', 68000, '2022-07-20', 'david.b@company.com'],
    [10, 'Maria Garcia', 'Marketing', 'Content Writer', 55000, '2023-09-01', 'maria.g@company.com'],
]
for row in employees:
    ws2.append(row)

for col in range(1, 8):
    ws2.column_dimensions[get_column_letter(col)].width = 20

# Sheet 3: Sales
ws3 = wb.create_sheet("Sales")
headers = ['Date', 'Product', 'Customer', 'Region', 'Quantity', 'Unit Price', 'Total']
ws3.append(headers)
for col in range(1, 8):
    ws3.cell(row=1, column=col).font = Font(bold=True)

products = ['Widget A', 'Widget B', 'Gadget Pro', 'Service Plan', 'Accessory Kit']
customers = ['Acme Corp', 'TechStart Inc', 'Global Solutions', 'MegaRetail', 'SmallBiz LLC']
regions = ['North', 'South', 'East', 'West']

random.seed(42)
for i in range(50):
    date = (datetime(2025, 1, 1) + timedelta(days=random.randint(0, 364))).strftime('%Y-%m-%d')
    product = random.choice(products)
    customer = random.choice(customers)
    region = random.choice(regions)
    qty = random.randint(1, 20)
    price = random.choice([25, 50, 100, 150, 200, 500])
    ws3.append([date, product, customer, region, qty, price, qty * price])

for col in range(1, 8):
    ws3.column_dimensions[get_column_letter(col)].width = 15

wb.save(os.path.join(TEST_DIR, 'financial_report.xlsx'))
print("Created: financial_report.xlsx")

# Create test PDF - Invoice
pdf_path = os.path.join(TEST_DIR, 'invoice.pdf')
c = canvas.Canvas(pdf_path, pagesize=letter)
width, height = letter

c.setFont("Helvetica-Bold", 24)
c.drawString(50, height - 50, "INVOICE")

c.setFont("Helvetica", 10)
c.drawString(50, height - 80, "From: Your Company LLC")
c.drawString(50, height - 95, "123 Business Street")
c.drawString(50, height - 110, "New York, NY 10001")
c.drawString(50, height - 125, "billing@company.com")

c.drawString(400, height - 80, "Invoice #: INV-2026-0142")
c.drawString(400, height - 95, "Date: January 15, 2026")
c.drawString(400, height - 110, "Due Date: February 15, 2026")

c.drawString(50, height - 160, "Bill To:")
c.drawString(50, height - 175, "Acme Corporation")
c.drawString(50, height - 190, "456 Client Avenue")
c.drawString(50, height - 205, "San Francisco, CA 94102")

c.setFont("Helvetica-Bold", 12)
c.drawString(50, height - 240, "Description")
c.drawString(300, height - 240, "Qty")
c.drawString(360, height - 240, "Unit Price")
c.drawString(450, height - 240, "Amount")

c.setStrokeColor(colors.grey)
c.line(50, height - 245, 550, height - 245)

c.setFont("Helvetica", 10)
y = height - 265
items = [
    ("Enterprise Software License (Annual)", 2, 5000),
    ("Premium Support Package", 1, 2500),
    ("Cloud Storage - 1TB", 12, 99),
    ("Custom Integration Development", 40, 150),
    ("Training Sessions (8 hours)", 2, 800),
]
total = 0
for desc, qty, price in items:
    amount = qty * price
    total += amount
    c.drawString(50, y, desc)
    c.drawString(300, y, str(qty))
    c.drawString(360, y, f"${price}")
    c.drawString(450, y, f"${amount}")
    y -= 15

c.line(50, y + 10, 550, y + 10)
c.setFont("Helvetica-Bold", 12)
c.drawString(350, y - 10, "TOTAL:")
c.drawString(450, y - 10, f"${total:,}")

c.setFont("Helvetica", 9)
c.drawString(50, y - 50, "Payment Terms: Net 30")
c.drawString(50, y - 65, "Payment Methods: Bank Transfer, Credit Card, PayPal")
c.drawString(50, y - 80, "Please include invoice number in payment reference")

c.setFont("Helvetica-Oblique", 8)
c.drawString(50, 30, "Thank you for your business!")

c.save()
print("Created: invoice.pdf")

# Create test CSV - 50 rows with demo data
import csv
with open(os.path.join(TEST_DIR, 'customers.csv'), 'w', newline='') as f:
    writer = csv.writer(f)
    writer.writerow(['ID', 'Name', 'Email', 'Phone', 'Amount', 'Date', 'Status'])
    
    names = ['John Smith', 'Sarah Johnson', 'Michael Chen', 'Emily Davis', 'Robert Wilson',
             'Lisa Anderson', 'James Taylor', 'Jennifer Lee', 'David Brown', 'Maria Garcia',
             'William Martinez', 'Emma Thompson', 'Daniel Jackson', 'Olivia White', 'Christopher Harris']
    domains = ['gmail.com', 'yahoo.com', 'outlook.com', 'company.com', 'business.org']
    statuses = ['Active', 'Pending', 'Completed', 'Cancelled']
    
    random.seed(123)
    for i in range(1, 51):
        name = random.choice(names)
        email = f"{name.lower().replace(' ', '.')}{i}@{random.choice(domains)}"
        phone = f"+1-{random.randint(200,999)}-{random.randint(100,999)}-{random.randint(1000,9999)}"
        amount = random.randint(100, 10000)
        date = (datetime(2025, 1, 1) + timedelta(days=random.randint(0, 364))).strftime('%Y-%m-%d')
        status = random.choice(statuses)
        writer.writerow([i, name, email, phone, amount, date, status])

print("Created: customers.csv")

# Create test TXT - Log file
log_entries = []
levels = ['INFO', 'DEBUG', 'WARNING', 'ERROR']
actions = ['User login', 'Data export', 'File upload', 'API request', 'Database query', 
           'Email sent', 'Payment processed', 'User created', 'Settings updated', 'Report generated']

random.seed(456)
for i in range(100):
    timestamp = (datetime(2026, 1, 14, 0, 0) + timedelta(minutes=i*15)).strftime('%Y-%m-%d %H:%M:%S')
    level = random.choice(levels)
    action = random.choice(actions)
    user = random.choice(['admin', 'user1', 'user2', 'system', 'api_service'])
    log_entries.append(f"[{timestamp}] [{level}] [{user}] {action}")

with open(os.path.join(TEST_DIR, 'app.log'), 'w') as f:
    f.write("=== Application Log File ===\n")
    f.write(f"Generated: 2026-01-14 23:59:59\n")
    f.write(f"Total Entries: {len(log_entries)}\n")
    f.write("=" * 50 + "\n\n")
    f.write("\n".join(log_entries))

print("Created: app.log")

print("\n=== All test files with demo data created ===")
