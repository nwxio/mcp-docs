#!/usr/bin/env python3
"""
MCP server for one-time file upload links
"""

import os
import sys
import json
import time
import requests
from datetime import datetime
from typing import List, Dict, Any, Optional
from urllib.parse import urlsplit, urlunsplit

from mcp.server import Server
from mcp.server.stdio import stdio_server
import mcp.types as types

UPLOADS_DIR = "/files/uploads"
TOKENS_FILE = "/files/tokens.json"
UPLOAD_SERVER_URL = "http://172.24.1.204:8765"
DOMAIN = "https://files.netwize.work"

os.makedirs(UPLOADS_DIR, exist_ok=True)

app = Server("file-upload")


def make_link(url: str, text: str = "Источник") -> str:
    return f"[{text}]({url})"


def rewrite_domain(raw_url: str, desired_domain: str) -> str:
    """
    Replace scheme+netloc with desired_domain, keep path/query/fragment intact.
    desired_domain example: "https://files.netwize.work"
    """
    if not raw_url:
        return raw_url

    src = urlsplit(raw_url)
    dst = urlsplit(desired_domain)

    # If raw_url is path-only
    if not src.scheme and raw_url.startswith("/"):
        return desired_domain.rstrip("/") + raw_url

    # If raw_url is missing scheme but has netloc-like string, keep as-is (rare)
    # urlsplit("example.com/x") treats as path; we won't guess here.
    if not src.scheme and not src.netloc and not raw_url.startswith("/"):
        return raw_url

    scheme = dst.scheme or src.scheme
    netloc = dst.netloc
    return urlunsplit((scheme, netloc, src.path, src.query, src.fragment))


def looks_like_single_b64_path(url: str) -> bool:
    """
    Heuristic: if path is a single very long segment (no slashes),
    it may be a fully-encoded blob path and could be a wrong kind of link.
    """
    try:
        u = urlsplit(url)
        path = (u.path or "").strip("/")
        if not path:
            return False
        segments = path.split("/")
        return len(segments) == 1 and len(segments[0]) > 60
    except Exception:
        return False


def create_upload_link(description: str = "") -> Dict:
    """Create a one-time upload link via API"""
    try:
        resp = requests.post(
            f"{UPLOAD_SERVER_URL}/api/create_token",
            json={"description": description},
            timeout=5
        )
        if resp.status_code == 200:
            return resp.json()
        return {'error': f"API error: {resp.status_code}"}
    except Exception as e:
        return {'error': str(e)}


def check_token(token: str) -> Dict:
    """Check if token has been used"""
    try:
        resp = requests.get(
            f"{UPLOAD_SERVER_URL}/api/check/{token}",
            timeout=5
        )
        if resp.status_code == 200:
            return resp.json()
        return {'exists': False, 'error': f"API error: {resp.status_code}"}
    except Exception as e:
        return {'exists': False, 'error': str(e)}


def process_file(filepath: str) -> Dict:
    """Auto-detect and process file"""
    if not os.path.exists(filepath):
        return {'success': False, 'error': 'File not found'}

    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    text_parts.append(f"--- Page {i+1} ---\n{text}")
                pages_count = len(pdf.pages)
            return {
                'success': True,
                'type': 'pdf',
                'pages': pages_count,
                'content': '\n\n'.join(text_parts)
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}

    elif ext in ['.docx', '.doc']:
        try:
            from docx import Document
            doc = Document(filepath)
            content = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                content.append("\n[TABLE]")
                for row in table.rows:
                    content.append(" | ".join(cell.text for cell in row.cells))
            return {'success': True, 'type': 'docx', 'content': '\n'.join(content)}
        except Exception as e:
            return {'success': False, 'error': str(e)}

    elif ext in ['.xlsx', '.xls']:
        try:
            from openpyxl import load_workbook
            wb = load_workbook(filepath, data_only=True)
            result = {'success': True, 'type': 'xlsx', 'sheets': {}}
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                data = []
                for row in ws.iter_rows(values_only=True):
                    data.append([str(c) if c is not None else "" for c in row])
                result['sheets'][sheet_name] = data
            return result
        except Exception as e:
            return {'success': False, 'error': str(e)}

    elif ext == '.csv':
        try:
            import pandas as pd
            df = pd.read_csv(filepath)
            return {
                'success': True,
                'type': 'csv',
                'rows': len(df),
                'columns': list(df.columns),
                'content': df.to_json(orient='records', force_ascii=False)
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}

    elif ext == '.json':
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return {'success': True, 'type': 'json', 'content': json.dumps(data, ensure_ascii=False, indent=2)}
        except Exception as e:
            return {'success': False, 'error': str(e)}

    elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp']:
        try:
            from PIL import Image
            with Image.open(filepath) as img:
                return {
                    'success': True,
                    'type': 'image',
                    'format': img.format,
                    'size': img.size,
                    'info': f"Image: {img.size[0]}x{img.size[1]}, {img.format}"
                }
        except Exception as e:
            return {'success': False, 'error': str(e)}

    else:
        # Try to read as text
        try:
            with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read(50000)
            return {'success': True, 'type': 'text', 'content': content, 'truncated': len(content) == 50000}
        except Exception as e:
            return {'success': False, 'error': str(e)}


@app.list_tools()
async def list_tools() -> List[types.Tool]:
    return [
        types.Tool(
            name="create_upload_link",
            description="Create a one-time upload link. User clicks link, uploads file, link becomes invalid. Returns URL for user.",
            inputSchema={
                "type": "object",
                "properties": {
                    "description": {"type": "string", "description": "Description for this upload (optional)"}
                },
                "required": []
            }
        ),
        types.Tool(
            name="check_upload",
            description="Check if user has uploaded a file via the link",
            inputSchema={
                "type": "object",
                "properties": {
                    "token": {"type": "string", "description": "Token from create_upload_link"}
                },
                "required": ["token"]
            }
        ),
        types.Tool(
            name="process_upload",
            description="Process the uploaded file (read content). File must be uploaded first.",
            inputSchema={
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Filename returned by check_upload"},
                    "delete_after": {"type": "boolean", "default": False, "description": "Delete file after processing"}
                },
                "required": ["filename"]
            }
        ),
        types.Tool(
            name="list_uploads",
            description="List all files in uploads directory",
            inputSchema={
                "type": "object",
                "properties": {},
                "required": []
            }
        ),
        types.Tool(
            name="delete_upload",
            description="Delete a file from uploads directory",
            inputSchema={
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Filename to delete"}
                },
                "required": ["filename"]
            }
        )
    ]


@app.call_tool()
async def call_tool(name: str, arguments: Dict[str, Any]) -> List[types.TextContent]:
    try:
        if name == "create_upload_link":
            description = arguments.get("description", "")
            result = create_upload_link(description)

            if 'error' in result:
                return [types.TextContent(type="text", text=f"Error creating link: {result['error']}")]

            raw_url = result.get('short_url') or result.get('url', '')
            url = rewrite_domain(raw_url, DOMAIN)

            # If URL looks suspicious, warn but still return link (safe fallback).
            warning = ""
            if looks_like_single_b64_path(url):
                warning = (
                    "\n\n⚠️ Warning: URL path looks like a single encoded blob (no slashes). "
                    "If downloads fail, check server-side link builder/shortener."
                )

            return [types.TextContent(
                type="text",
                text=f"{make_link(url, 'Загрузить')}{warning}\n\n⏱️ Link expires in 30 minutes"
            )]

        elif name == "check_upload":
            token = arguments["token"]
            result = check_token(token)

            if 'error' in result and not result.get('exists'):
                return [types.TextContent(type="text", text=f"Token error: {result.get('error', 'Invalid token')}")]

            if result.get('expired'):
                return [types.TextContent(type="text", text="❌ Upload link has expired")]

            if result.get('used'):
                filename = result.get('filename')
                if not filename:
                    return [types.TextContent(type="text", text="Error: upload marked as used but filename is missing")]
                filepath = os.path.join(UPLOADS_DIR, filename)
                size = os.path.getsize(filepath) if os.path.exists(filepath) else 0
                return [types.TextContent(
                    type="text",
                    text=f"✅ File uploaded!\n\n"
                         f"Filename: {filename}\n"
                         f"Size: {size / 1024:.1f} KB\n\n"
                         f"Use `process_upload` to read the file."
                )]
            else:
                return [types.TextContent(type="text", text="⏳ Waiting for upload... File not uploaded yet.")]

        elif name == "process_upload":
            filename = arguments["filename"]
            delete_after = arguments.get("delete_after", False)
            filepath = os.path.join(UPLOADS_DIR, filename)

            if not os.path.exists(filepath):
                return [types.TextContent(type="text", text=f"File not found: {filename}")]

            result = process_file(filepath)

            if result.get('success'):
                content = result.get('content', '')

                # For structured data
                if result.get('type') == 'xlsx':
                    output = f"📁 File: {filename}\n📊 Type: XLSX\n\n{json.dumps(result['sheets'], ensure_ascii=False, indent=2)}"
                elif result.get('type') == 'csv':
                    output = f"📁 File: {filename}\n📊 Type: CSV\nRows: {result['rows']}\nColumns: {result['columns']}\n\n{content[:10000]}"
                elif result.get('type') == 'image':
                    output = f"📁 File: {filename}\n🖼️ Type: Image\n{result['info']}"
                else:
                    if isinstance(content, str) and len(content) > 30000:
                        content = content[:30000] + "\n\n... [content truncated]"
                    output = f"📁 File: {filename}\n📄 Type: {result.get('type', 'unknown')}\n\n{content}"

                if delete_after:
                    os.remove(filepath)
                    output += "\n\n🗑️ File deleted after processing."

                return [types.TextContent(type="text", text=output)]
            else:
                return [types.TextContent(type="text", text=f"Error processing {filename}: {result.get('error', 'Unknown error')}")]

        elif name == "list_uploads":
            files = []
            for f in os.listdir(UPLOADS_DIR):
                filepath = os.path.join(UPLOADS_DIR, f)
                if os.path.isfile(filepath):
                    stat = os.stat(filepath)
                    files.append({
                        'name': f,
                        'size': stat.st_size,
                        'modified': datetime.fromtimestamp(stat.st_mtime).isoformat()[:19]
                    })

            if not files:
                return [types.TextContent(type="text", text="No files in uploads directory.")]

            files.sort(key=lambda x: x['modified'], reverse=True)
            output = f"Files in uploads ({len(files)}):\n\n"
            for f in files:
                size_str = f"{f['size'] / 1024:.1f} KB" if f['size'] < 1024 * 1024 else f"{f['size'] / (1024 * 1024):.1f} MB"
                output += f"• {f['name']} ({size_str}) - {f['modified']}\n"

            return [types.TextContent(type="text", text=output)]

        elif name == "delete_upload":
            filename = arguments["filename"]
            filepath = os.path.join(UPLOADS_DIR, filename)

            if not os.path.exists(filepath):
                return [types.TextContent(type="text", text=f"File not found: {filename}")]

            os.remove(filepath)
            return [types.TextContent(type="text", text=f"🗑️ Deleted: {filename}")]

        else:
            return [types.TextContent(type="text", text=f"Unknown tool: {name}")]

    except Exception as e:
        return [types.TextContent(type="text", text=f"Error: {str(e)}")]


async def main():
    async with stdio_server() as (read_stream, write_stream):
        await app.run(read_stream, write_stream, app.create_initialization_options())


if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
